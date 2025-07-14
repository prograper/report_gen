"""
入口脚本：解析 Excel → 结构化抽取 → 生成段落 → 写入 Word
兼容任何遵循 OpenAI Chat Completion & Function-Calling 协议的 LLM。
"""

import os, sys, yaml, pandas as pd
from pathlib import Path
from docx import Document
from agents.registry import get_extractor
from agents.generate import get_generator

# ───── 目录常量 ───────────────────────────────────────────
ROOT       = Path(__file__).parent
# CFG_DIR    = ROOT / "configs"
CFG_DIR   = ROOT
PROMPT_DIR = ROOT / "prompts"
TPL_DIR    = ROOT / "templates"

# ───── 实用函数 ───────────────────────────────────────────
def load_yaml(cfg_dir: str, fname: str):
    with open(cfg_dir / "business_configs" / fname, encoding="utf-8") as f:
        return yaml.safe_load(f)

def read_prompt(path: str):
    return (ROOT / path).read_text(encoding="utf-8")

def write_docx(config_dir: str, report_name:str, placeholder_map: dict, paragraph_map: dict):
    doc = Document(config_dir / "template" / "report_template.docx")
    for para_id, texts in paragraph_map.items():
        placeholder = "{{" + placeholder_map[para_id] + "}}"
        for p in doc.paragraphs:
            if placeholder in p.text:
                p.text = p.text.replace(placeholder, "")
                for seg in texts:
                    run = p.add_run(seg)
                    run.add_break()
                break
    doc.save(config_dir / "output" / (report_name + ".docx"))

# -------- 路径解析工具 --------
def resolve(path: str, data: dict):
    cur = data
    for part in path.split("."):
        if isinstance(cur, dict) and part in cur:
            cur = cur[part]
        else:
            return None
    return cur

# 找到/config/business_configs/input下的第一个 Excel 文件
def first_excel_as_excelfile(
        dir_path: str | Path,
        pattern: str = "*.xls*"
    ) -> pd.ExcelFile:
    """
    在 dir_path 中按 pattern 找到第一个 Excel 文件，
    并以 pd.ExcelFile 形式返回。
    """
    dir_path = Path(dir_path)
    matches = sorted(dir_path.glob(pattern))   # 如需递归改成 rglob

    if not matches:
        raise FileNotFoundError(
            f"目录 {dir_path} 下没有找到任何 {pattern} 文件！")
    
    if len(matches) > 1:
        # 如果确实只会放一个文件，可以改成直接 matches[0] 而不提示
        print(f"⚠️  注意：发现 {len(matches)} 个 Excel，仅使用第一个：{matches[0].name}")

    return pd.ExcelFile(matches[0])

# ───── Pipeline ─────────────────────────────────────────
def run_pipeline(config_dir: str, report_name: str):
    config_dir = Path(config_dir)
    # ───── 读取配置 ───────────────────────────────────────────
    # 每张 Sheet 抽取任务
    sheet_cfg   = load_yaml(config_dir, "sheet_tasks.yaml")
    # 段落生成任务
    paragraphs  = load_yaml(config_dir, "paragraph_tasks.yaml")
    placeholder_map = load_yaml(config_dir, "doc_placeholders.yaml")
    
    xls = first_excel_as_excelfile(config_dir / "input", "*.xls*")
    # 全局嵌套 {Sheet: {field: val}}
    extracted = {}
    # 1) 遍历 Sheet 抽取
    for sheet in xls.sheet_names:
        if sheet not in sheet_cfg:
            continue
        cfg     = sheet_cfg[sheet]
        df      = xls.parse(sheet)
        extractor = get_extractor("GenericExtractor")(
            df          = df,
            keys        = cfg["keys"],
            prompt_path = config_dir / "prompts" / cfg["prompt"],
            config_dir  = config_dir,
            provider = "qwen"
        )
        # extracted.update(extractor.extract())
        # extractor.extract() 仍返回 {"Tmax":..., "Cmax":...}
        extracted[sheet] = extractor.extract()

    # 2) 生成段落
    para_out = {pid: [] for pid in paragraphs}
    for pid, task in paragraphs.items():
        # 取所需字段
        # ctx = {k: extracted.get(k) for k in task["keys"]}
        # if None in ctx.values():
        #     print(f"⚠️  段落 {pid} 缺字段，已跳过")
        #     continue
        # generator = get_generator("GenericParagraphGenerator")(
        #     prompt_path = task["prompt"],
        #     context     = ctx,

        # 只做存在性检查；模板里直接通过 data.<Sheet>.<Field> 访问
        missing = [k for k in task["keys"] if resolve(k, extracted) is None]
        if missing:
            print(f"⚠️  段落 {pid} 缺字段 {missing}，已跳过")
            continue
        generator = get_generator("GenericParagraphGenerator")(
            prompt_path = config_dir / "prompts" / task["prompt"],
            # ← 只放一层 data，如果在prompt里不想加data前缀，可以context = extracted
            context     = {"data": extracted},
            config_dir = config_dir,
            provider = "qwen"
        )
        para_out[pid].append(generator.generate())

    # 3) 写入 Word
    write_docx(config_dir, report_name, placeholder_map, para_out)
    print(f"✓ 已生成报告：{config_dir / 'output' / (report_name + '.docx')}")

# ───── CLI ──────────────────────────────────────────────
if __name__ == "__main__":
    if "DASHSCOPE_API_KEY" not in os.environ:
        sys.exit("✗ 请先 set DASHSCOPE_API_KEY=sk-...")
    import argparse
    ap = argparse.ArgumentParser(description="自动生成报告流水线")
    # ap.add_argument("excel", help="原始 Excel 路径")
    # ap.add_argument("-o", "--out", default="report_out.docx", help="输出 Word 文件名")
    ap.add_argument("-c", "--config", default="configs", help="配置文件目录")
    ap.add_argument("-n", "--name", default="生成报告文件", help="报告名称")
    args = ap.parse_args()

    run_pipeline(args.config, args.name)
